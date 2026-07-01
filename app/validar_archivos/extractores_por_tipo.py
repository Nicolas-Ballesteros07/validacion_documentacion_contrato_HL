"""
extractores_por_tipo.py
-----------------------
Extracción de campos con conocimiento de la estructura real de cada documento.
Reemplaza la búsqueda genérica por regex en texto plano, que falla porque
los datos están embebidos en prosa o en celdas de tabla con formatos propios
de cada tipo de documento.

Uso desde views.py:
    from .extractores_por_tipo import extraer_campos_estructurados
    campos_extraidos = extraer_campos_estructurados(tipo_archivo, archivo)
"""

import re
from docx import Document

# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

MESES_ES = {
    'enero': '01', 'febrero': '02', 'marzo': '03', 'abril': '04',
    'mayo': '05', 'junio': '06', 'julio': '07', 'agosto': '08',
    'septiembre': '09', 'octubre': '10', 'noviembre': '11', 'diciembre': '12',
}

PATRON_FECHA_LARGA = re.compile(
    r'(\d{1,2})\s+de\s+(\w+)\s+de\s+(\d{4})', re.IGNORECASE
)
PATRON_FECHA_CORTA = re.compile(r'(\d{1,2})[\/\-](\d{1,2})[\/\-](\d{4})')
PATRON_FECHA_ISO   = re.compile(r'(\d{4})[\/\-](\d{2})[\/\-](\d{2})')
PATRON_CC = re.compile(
    r'C\.?\s*C\.?\s*(?:No\.?)?\s*(\d[\d.\s]{4,}\d)',
    re.IGNORECASE
)
PATRON_MONTO = re.compile(r'\$\s*([\d.,]+)')

# Un "encabezado de sección" en estos documentos es una línea corta, casi
# toda en mayúsculas y sin ":" de campo (p.ej. "AUXILIO EXTRALEGAL DE VIVIENDA").
# Se usa como frontera para saber dónde termina una sección al buscar montos.
_PATRON_POSIBLE_ENCABEZADO = re.compile(r'^[A-ZÁÉÍÓÚÑ0-9º°.\-,\s]{8,90}$')


def _es_encabezado_seccion(texto):
    """Heurística para detectar si un párrafo es un título de sección
    (ej. 'AUXILIO EXTRALEGAL DE TRANSPORTE EN OBRA') y no un párrafo de
    contenido normal."""
    if not texto or len(texto) >= 90:
        return False
    return bool(_PATRON_POSIBLE_ENCABEZADO.match(texto))


def _normalizar_fecha_texto(texto):
    """Convierte cualquier formato de fecha encontrado a YYYY-MM-DD."""
    m = PATRON_FECHA_LARGA.search(texto)
    if m:
        dia, mes_str, anio = m.groups()
        mes = MESES_ES.get(mes_str.lower())
        if mes:
            return f"{anio}-{mes}-{dia.zfill(2)}"
    m = PATRON_FECHA_ISO.search(texto)
    if m:
        return f"{m.group(1)}-{m.group(2)}-{m.group(3)}"
    m = PATRON_FECHA_CORTA.search(texto)
    if m:
        d, mo, y = m.groups()
        # 06/17/2026 → detectar formato MM/DD/YYYY (usado en RIESGO)
        if int(d) > 12:        # primer número es día → DD/MM/YYYY
            return f"{y}-{mo.zfill(2)}-{d.zfill(2)}"
        elif int(mo) > 12:     # segundo número es día → MM/DD/YYYY
            return f"{y}-{d.zfill(2)}-{mo.zfill(2)}"
        else:                  # ambiguo → asumir DD/MM/YYYY
            return f"{y}-{mo.zfill(2)}-{d.zfill(2)}"
    return None


def _limpiar_cc(texto):
    """Extrae solo dígitos de un número de cédula."""
    return re.sub(r'\D', '', str(texto))


def _celda(tabla, fila, col):
    """Lee una celda de tabla de forma segura."""
    try:
        return tabla.rows[fila].cells[col].text.strip()
    except IndexError:
        return ""


def _parrafos(doc):
    return [p.text.strip() for p in doc.paragraphs if p.text.strip()]


def _extraer_monto_seccion(parrafos, patron_inicio, max_parrafos=15, patron_exclusion=None):
    """
    Busca un monto con formato "($1.234.567)" dentro de la sección que
    arranca en el primer párrafo que matchea `patron_inicio`.

    La sección se considera cerrada al llegar al siguiente encabezado de
    sección en mayúsculas (heurística `_es_encabezado_seccion`) o, como
    salvaguarda, tras `max_parrafos` párrafos sin encontrar nada.

    Esto reemplaza los regex anteriores que exigían una redacción exacta
    tipo "valor diario de PESOS COLOMBIANOS (...)". Esos regex fallaban en
    la práctica porque casi todos estos contratos escriben el monto en
    letras ENTRE la etiqueta y el paréntesis con el número
    (ej. "valor diario de VEINTISIETE MIL... PESOS COLOMBIANOS ($27.921)"),
    y a veces el monto queda en un párrafo distinto al de la etiqueta.
    Buscar el monto en toda la sección (no en una frase exacta) es mucho
    más tolerante a variaciones de redacción entre documentos.

    Si el mismo encabezado aparece más de una vez en el documento (poco
    común, pero posible en cartas con anexos repetidos), se sigue
    intentando con la siguiente ocurrencia hasta encontrar un monto.
    """
    for idx, p in enumerate(parrafos):
        if not patron_inicio.search(p):
            continue
        if patron_exclusion and patron_exclusion.search(p):
            continue

        for offset in range(max_parrafos):
            j = idx + offset
            if j >= len(parrafos):
                break
            # No cortar en el propio párrafo de encabezado (offset 0)
            if offset > 0 and _es_encabezado_seccion(parrafos[j]):
                break
            m = re.search(r'\(\s*\$\s*([\d.,]+)\s*\)', parrafos[j])
            if m:
                return re.sub(r'[.,]', '', m.group(1))
        # No se encontró monto en esta ocurrencia; seguir buscando por si
        # el encabezado se repite más adelante en el documento.
    return None


# ---------------------------------------------------------------------------
# Extractor AUXILIO
# ---------------------------------------------------------------------------

def _extraer_auxilio(doc):
    """
    Extrae todos los campos del documento "Vinculación-Auxilios":
      - nombre / documento / ciudad_expedicion: tabla de firma (C.C. NUMERO Expedida en CIUDAD)
      - cargo / proyecto: párrafo "desempeñándose como CARGO en el área PROYECTO"
      - fecha_inicio: párrafo de fecha de firma del documento
      - auxilio_localizacion: párrafo SEGUNDO del Pacto de Exclusión Salarial
        "reconocerá a EL TRABAJADOR la suma de … ($MONTO) como un auxilio extralegal
         de vivienda y manutención … será en CIUDAD"
      - ciudad: ciudad destino extraída del mismo párrafo SEGUNDO
      - aux_almuerzo: sección "AUXILIO EXTRALEGAL DE ALIMENTACIÓN EN OBRA"
      - aux_transporte: sección "AUXILIO EXTRALEGAL DE TRANSPORTE EN OBRA"
      - aux_vivienda: sección "AUXILIO EXTRALEGAL DE VIVIENDA" (si existe separada
        de "VIVIENDA Y MANUTENCION", que se mapea a auxilio_localizacion)
      - beneficio_alimentacion (Peoplepass): sección con "PEOPLEPASS"
      - aux_desplazamiento: sección "AUXILIO DE DESPLAZAMIENTO"
      - beneficio_pension: sección "APORTES VOLUNTARIOS A (FONDO DE) PENSIÓN(ES)"

    Todos los campos monetarios usan `_extraer_monto_seccion`, que busca el
    monto en toda la sección (no en una única frase exacta), por lo que es
    robusto ante variaciones de redacción entre plantillas/documentos.
    """
    campos = {}
    parrafos = _parrafos(doc)

    # ── nombre, documento y ciudad_expedicion desde tabla de firma ──────────
    for tabla in doc.tables:
        for fila in tabla.rows:
            for celda in fila.cells:
                texto_celda = celda.text.strip()
                if not texto_celda:
                    continue
                m_cc = PATRON_CC.search(texto_celda)
                if m_cc and 'nombre' not in campos:
                    lineas = [l.strip() for l in texto_celda.split('\n') if l.strip()]
                    if lineas:
                        campos['nombre'] = lineas[0]
                    campos['documento'] = _limpiar_cc(m_cc.group(1))
                    # ciudad de expedición: "Expedida en CIUDAD" o "Expedida CIUDAD"
                    m_ciudad_exp = re.search(
                        r'[Ee]xpedida\s+(?:en\s+)?([A-ZÁÉÍÓÚÑ][A-ZÁÉÍÓÚÑA-Za-záéíóúña-z\s]+?)(?:\s*$|\n)',
                        texto_celda
                    )
                    if m_ciudad_exp:
                        campos['ciudad_expedicion'] = m_ciudad_exp.group(1).strip()
                if 'nombre' in campos and 'documento' in campos:
                    break
            if 'nombre' in campos and 'documento' in campos:
                break

    # ── cargo y proyecto desde párrafo introductorio ────────────────────────
    pat_cargo_proy = re.compile(
        r'desempe[ñn][aá]ndose\s+como\s+(.+?)\s+en\s+el\s+[aá]rea\s+(.+)',
        re.IGNORECASE
    )
    for p in parrafos:
        m = pat_cargo_proy.search(p)
        if m:
            campos['cargo']    = m.group(1).strip()
            campos['proyecto'] = m.group(2).strip()
            break

    # ── fecha_inicio: última "el DD de mes de YYYY" en el doc ───────────────
    pat_firma = re.compile(r'\bel\s+(\d{1,2}\s+de\s+\w+\s+de\s+\d{4})', re.IGNORECASE)
    for p in reversed(parrafos):
        m = pat_firma.search(p)
        if m:
            fn = _normalizar_fecha_texto(m.group(1))
            if fn:
                campos['fecha_inicio'] = fn
                break

    # ── auxilio_localizacion y ciudad destino ────────────────────────────────
    # Párrafo SEGUNDO: "reconocerá a EL TRABAJADOR la suma de … ($MONTO)
    #  … será en CIUDAD."
    pat_segundo = re.compile(
        r'SEGUNDO[.\s]+AUXILIO EXTRALEGAL DE VIVIENDA',
        re.IGNORECASE
    )
    for p in parrafos:
        if pat_segundo.search(p):
            m_ciudad = re.search(
                r'ser[aá]\s+en\s+([A-ZÁÉÍÓÚÑA-Za-záéíóúña-z\s]+?)[\.\s]*$',
                p
            )
            if m_ciudad:
                campos['ciudad'] = m_ciudad.group(1).strip().rstrip('.')
            break
    v = _extraer_monto_seccion(parrafos, pat_segundo, max_parrafos=3)
    if v:
        campos['auxilio_localizacion'] = v

    # ── aux_almuerzo: AUXILIO EXTRALEGAL DE ALIMENTACIÓN EN OBRA ────────────
    pat_almuerzo = re.compile(r'AUXILIO EXTRALEGAL DE ALIMENTACI[OÓ]N EN OBRA', re.IGNORECASE)
    v = _extraer_monto_seccion(parrafos, pat_almuerzo)
    if v:
        campos['aux_almuerzo'] = v

    # ── aux_transporte: AUXILIO EXTRALEGAL DE TRANSPORTE EN OBRA ────────────
    pat_transporte = re.compile(r'AUXILIO EXTRALEGAL DE TRANSPORTE EN OBRA', re.IGNORECASE)
    v = _extraer_monto_seccion(parrafos, pat_transporte)
    if v:
        campos['aux_transporte'] = v

    # ── aux_vivienda: AUXILIO EXTRALEGAL DE VIVIENDA (sección propia) ───────
    # Se excluye el encabezado combinado "VIVIENDA Y MANUTENCION", que ya se
    # captura como auxilio_localizacion más arriba.
    pat_vivienda = re.compile(r'AUXILIO EXTRALEGAL DE VIVIENDA', re.IGNORECASE)
    pat_excl_vivienda = re.compile(r'VIVIENDA\s+Y\s+MANUTENCION', re.IGNORECASE)
    v = _extraer_monto_seccion(parrafos, pat_vivienda, patron_exclusion=pat_excl_vivienda)
    if v:
        campos['aux_vivienda'] = v

    # ── beneficio_alimentacion (Peoplepass) ─────────────────────────────────
    pat_peoplepass = re.compile(r'PEOPLEPASS', re.IGNORECASE)
    v = _extraer_monto_seccion(parrafos, pat_peoplepass)
    if v:
        campos['beneficio_alimentacion'] = v

    # ── aux_desplazamiento: sección "AUXILIO DE DESPLAZAMIENTO" ─────────────
    pat_desplazamiento = re.compile(r'^AUXILIO DE DESPLAZAMIENTO\s*$', re.IGNORECASE)
    v = _extraer_monto_seccion(parrafos, pat_desplazamiento)
    if v:
        campos['aux_desplazamiento'] = v

    # ── beneficio_pension: BENEFICIO EXTRALEGAL APORTE VOLUNTARIO A PENSIONES
    pat_pension = re.compile(
        r'APORTES?\s+VOLUNTARIOS?\s+A\s+(?:FONDO\s+DE\s+)?PENSI[OÓ]N(?:ES)?',
        re.IGNORECASE
    )
    v = _extraer_monto_seccion(parrafos, pat_pension)
    if v:
        campos['beneficio_pension'] = v

    return campos


# ---------------------------------------------------------------------------
# Extractor CONDICIONES
# ---------------------------------------------------------------------------

def _extraer_condiciones(doc):
    """
    Carta de condiciones laborales con datos inline:
      "Cargo: VALOR"
      "Salario en Colombia mensual: $VALOR"
      "Bonos de Alimentación ... mensuales: $VALOR"
      "Auxilio de Localización: $VALOR"
      "Pensión Voluntaria Institucional mensual: $VALOR"
      "Descanso: VALOR"
    """
    campos = {}
    parrafos = _parrafos(doc)

    patrones_inline = {
        'cargo':                  re.compile(r'^Cargo\s*:\s*(.+)$', re.IGNORECASE),
        'salario_basico':         re.compile(r'Salario en Colombia mensual\s*:\s*\$([\d.,]+)', re.IGNORECASE),
        'beneficio_alimentacion': re.compile(r'Bonos?\s+de\s+Alimentaci[oó]n[^:]*:\s*\$([\d.,]+)', re.IGNORECASE),
        'auxilio_localizacion':   re.compile(r'Auxilio de Localizaci[oó]n\s*:\s*\$([\d.,]+)', re.IGNORECASE),
        'beneficio_pension':      re.compile(r'Pensi[oó]n Voluntaria[^:]*:\s*\$([\d.,]+)', re.IGNORECASE),
        'condicion_descanso':     re.compile(r'^Descanso\s*:\s*(.+)$', re.IGNORECASE),
    }

    for p in parrafos:
        # Fecha de carta (primer párrafo)
        if not campos.get('fecha_inicio'):
            fn = _normalizar_fecha_texto(p)
            if fn:
                campos['fecha_inicio'] = fn

        # Nombre: primera línea en mayúsculas puras (no es etiqueta)
        if not campos.get('nombre'):
            if re.match(r'^[A-ZÁÉÍÓÚÑ][A-ZÁÉÍÓÚÑ\s]{4,}$', p) and ':' not in p:
                # Excluir etiquetas conocidas
                if not any(x in p.upper() for x in ('SEÑOR', 'ESTIMADO', 'HL INFRA')):
                    campos['nombre'] = p

        for clave, patron in patrones_inline.items():
            if clave not in campos:
                m = patron.search(p)
                if m:
                    campos[clave] = m.group(1).strip()

    # proyecto: buscar "proyecto CODIGO NOMBRE" en el párrafo introductorio
    pat_proy = re.compile(r'proyecto\s+([\w\d]+\s+[\w\d\s]+?)(?:[,.]|$)', re.IGNORECASE)
    for p in parrafos:
        m = pat_proy.search(p)
        if m and not campos.get('proyecto'):
            campos['proyecto'] = m.group(1).strip()

    return campos


# ---------------------------------------------------------------------------
# Extractor CONTRATO
# ---------------------------------------------------------------------------

def _extraer_contrato(doc):
    """
    Tabla de encabezado del contrato:
      fila 3 col 0: nombre
      fila 5 col 0: CC | fila 5 col 1: teléfono
      fila 7 col 1: cargo
      fila 9 col 0: salario (letras + cifra)
      fila 11 col 0: fecha
    """
    campos = {}
    if not doc.tables:
        return campos

    t = doc.tables[0]

    nombre = _celda(t, 3, 0)
    if nombre:
        campos['nombre'] = nombre

    cc_raw = _celda(t, 5, 0)
    if cc_raw:
        nums = [n for n in re.findall(r'\d+', cc_raw) if len(n) >= 6]
        if nums:
            campos['documento'] = max(nums, key=len)

    tel_raw = _celda(t, 5, 1)
    if tel_raw:
        campos['numero_contacto'] = re.sub(r'\D', '', tel_raw)

    cargo = _celda(t, 7, 1)
    if cargo:
        campos['cargo'] = cargo

    salario_raw = _celda(t, 9, 0)
    if salario_raw:
        m = PATRON_MONTO.search(salario_raw)
        if m:
            # Normalizar quitando separadores de miles → entero como string
            campos['salario_basico'] = re.sub(r'[.,]', '', m.group(1))

    fecha_raw = _celda(t, 11, 0)
    if fecha_raw:
        fn = _normalizar_fecha_texto(fecha_raw)
        if fn:
            campos['fecha_inicio'] = fn

    return campos


# ---------------------------------------------------------------------------
# Extractor MAN_CONF
# ---------------------------------------------------------------------------

def _extraer_man_conf(doc):
    """
    Otrosí manual de confidencialidad:
      Tabla 0 fila 0 celda 1:
        "EL TRABAJADOR\n...\nNombre:  NOMBRE\nC.C. No. NUMERO de CIUDAD"
      Fecha en párrafo de cierre.
    """
    campos = {}
    parrafos = _parrafos(doc)

    if doc.tables:
        celda = _celda(doc.tables[0], 0, 1)

        m_nombre = re.search(r'Nombre\s*:\s*(.+)', celda, re.IGNORECASE)
        if m_nombre:
            campos['nombre'] = m_nombre.group(1).strip()
        else:
            # fallback: línea en mayúsculas que no sea "EL TRABAJADOR"
            for linea in celda.split('\n'):
                linea = linea.strip()
                if re.match(r'^[A-ZÁÉÍÓÚÑ][A-ZÁÉÍÓÚÑ\s]{4,}$', linea):
                    if 'TRABAJADOR' not in linea.upper():
                        campos['nombre'] = linea
                        break

        m_cc = PATRON_CC.search(celda)
        if m_cc:
            campos['documento'] = _limpiar_cc(m_cc.group(1))

    # fecha desde el párrafo de firma
    pat_firma = re.compile(r'\bel\s+(\d{1,2}\s+de\s+\w+\s+de\s+\d{4})', re.IGNORECASE)
    for p in parrafos:
        m = pat_firma.search(p)
        if m:
            fn = _normalizar_fecha_texto(m.group(1))
            if fn:
                campos['fecha_inicio'] = fn
                break

    return campos


# ---------------------------------------------------------------------------
# Extractor MEDICAS
# ---------------------------------------------------------------------------

def _extraer_medicas(doc):
    """
    Carta de resultados médicos:
      - fecha_inicio: primer párrafo (fecha de la carta)
      - nombre: párrafo en mayúsculas justo después de "Señor (a)"
        (puede estar en orden apellido-nombre o nombre-apellido)
      - cargo: párrafo "contratado para el cargo de CARGO" — solo el cargo,
        cortando antes de cualquier texto adicional
    """
    campos = {}
    parrafos = _parrafos(doc)

    # fecha de carta
    for p in parrafos[:4]:
        fn = _normalizar_fecha_texto(p)
        if fn:
            campos['fecha_inicio'] = fn
            break

    # nombre: primera línea en mayúsculas puras (excluir etiquetas)
    EXCLUIR = {'SEÑOR (A)', 'SEÑOR(A)', 'E.S.M.', 'HL INFRAESTRUCTURA S.A.S.'}
    for p in parrafos:
        if re.match(r'^[A-ZÁÉÍÓÚÑ][A-ZÁÉÍÓÚÑ\s]{4,}$', p) and p.upper() not in EXCLUIR:
            campos['nombre'] = p
            break

    # cargo: "contratado para el cargo de CARGO" — cortar en espacio/punto/coma
    pat_cargo = re.compile(
        r'contratado\s+para\s+el\s+cargo\s+de\s+([A-ZÁÉÍÓÚ][A-ZÁÉÍÓÚ\s]+?)(?:\s+(?:cualquier|en|y|,|\.)|$)',
        re.IGNORECASE
    )
    for p in parrafos:
        m = pat_cargo.search(p)
        if m:
            campos['cargo'] = m.group(1).strip().rstrip('.,')
            break

    return campos


# ---------------------------------------------------------------------------
# Extractor RIESGO
# ---------------------------------------------------------------------------
def _detectar_tipo_riesgo(doc):
    """
    Determina si el documento de Riesgos corresponde a 'Obra' u 'Oficina'
    según si hay alguna 'X' marcada en las columnas de casilla de la tabla
    de peligros (tabla índice 1: 'PELIGROS A LOS QUE ESTÁ EXPUESTO').

    Estructura de la tabla: 6 columnas por fila.
      col 0/1: nombre del peligro (columna izquierda, fusionada)
      col 2:   casilla de marcado para el peligro de la izquierda
      col 3/4: nombre del peligro (columna derecha, fusionada)
      col 5:   casilla de marcado para el peligro de la derecha

    Si CUALQUIER casilla (col 2 o col 5) tiene contenido no vacío en
    alguna fila de peligro real (se excluyen encabezados y la fila de
    firma), el documento es de 'OFICINA'. Si ninguna casilla tiene
    contenido, es de 'OBRA'.
    """
    if len(doc.tables) < 2:
        return 'OBRA'  # sin tabla de peligros, no se puede determinar; default seguro

    tabla_peligros = doc.tables[1]

    # Etiquetas de fila que no son peligros reales (encabezados/firma) y
    # deben ignorarse aunque su celda de "casilla" tenga texto fusionado.
    FILAS_EXCLUIDAS = {
        'PELIGROS A LOS QUE ESTÁ EXPUESTO',
        'FIRMA TRABAJADOR:',
    }

    for row in tabla_peligros.rows:
        cells = row.cells
        if len(cells) < 6:
            continue

        etiqueta_izq = cells[0].text.strip()
        etiqueta_der = cells[3].text.strip()

        casilla_izq = cells[2].text.strip()
        casilla_der = cells[5].text.strip()

        es_fila_excluida_izq = any(etiqueta_izq.startswith(f) for f in FILAS_EXCLUIDAS)
        es_fila_excluida_der = any(etiqueta_der.startswith(f) for f in FILAS_EXCLUIDAS)

        if casilla_izq and not es_fila_excluida_izq:
            return 'OFICINA'
        if casilla_der and not es_fila_excluida_der:
            return 'OFICINA'

    return 'OBRA'



def _extraer_riesgo(doc):
    """
    Notificación de riesgos — tabla 0, fila 0:
      celda 0: "NOMBRE\nNombre\n\nCC Expedida en CIUDAD\nNúmero de documento\n\nCARGO\nCargo"
      celda 2: "PROYECTO\nCentro de trabajo / Proyecto\n\nFECHA\nFecha de ingreso"
    """
    campos = {}
    if not doc.tables:
        return campos

    celda0 = _celda(doc.tables[0], 0, 0)
    celda2 = _celda(doc.tables[0], 0, 2)

    def _extraer_antes_de_etiqueta(texto, etiqueta_re):
        lineas = [l.strip() for l in texto.split('\n') if l.strip()]
        for i, linea in enumerate(lineas):
            if etiqueta_re.match(linea):
                if i > 0:
                    return lineas[i - 1]
        return None

    # nombre: línea antes de la etiqueta "Nombre"
    v = _extraer_antes_de_etiqueta(celda0, re.compile(r'^nombre$', re.IGNORECASE))
    if v:
        campos['nombre'] = v

    # documento: línea antes de "Número de documento"
    v = _extraer_antes_de_etiqueta(celda0, re.compile(r'^n[úu]mero\s+de\s+documento$', re.IGNORECASE))
    if v:
        campos['documento'] = _limpiar_cc(v)

    # cargo: línea antes de "Cargo"
    v = _extraer_antes_de_etiqueta(celda0, re.compile(r'^cargo$', re.IGNORECASE))
    if v:
        campos['cargo'] = v

    # proyecto: línea antes de "Centro de trabajo / Proyecto"
    v = _extraer_antes_de_etiqueta(celda2, re.compile(r'^centro\s+de\s+trabajo', re.IGNORECASE))
    if v:
        campos['proyecto'] = v

    # fecha: línea antes de "Fecha de ingreso"
    v = _extraer_antes_de_etiqueta(celda2, re.compile(r'^fecha\s+de\s+ingreso$', re.IGNORECASE))
    if v:
        fn = _normalizar_fecha_texto(v)
        if fn:
            campos['fecha_inicio'] = fn

    campos['tipo_riesgo'] = _detectar_tipo_riesgo(doc)

    return campos


# ---------------------------------------------------------------------------
# Dispatcher principal
# ---------------------------------------------------------------------------

EXTRACTORES = {
    'AUXILIO':     _extraer_auxilio,
    'CONDICIONES': _extraer_condiciones,
    'CONTRATO':    _extraer_contrato,
    'MAN_CONF':    _extraer_man_conf,
    'MEDICAS':     _extraer_medicas,
    'RIESGO':      _extraer_riesgo,
}


def extraer_campos_estructurados(tipo_documento, archivo):
    """
    Extrae los campos del documento usando la estructura conocida de cada tipo.

    Args:
        tipo_documento: str — 'AUXILIO', 'CONDICIONES', 'CONTRATO',
                                'MAN_CONF', 'MEDICAS', 'RIESGO'
        archivo: objeto file-like (request.FILES item) o path str

    Returns:
        dict con las claves disponibles para ese tipo:
          nombre, documento, cargo, proyecto, fecha_inicio (formato YYYY-MM-DD),
          salario_basico, beneficio_alimentacion, auxilio_localizacion,
          beneficio_pension, condicion_descanso, numero_contacto
        Todos los valores son strings. Las claves ausentes simplemente no
        aparecen en el dict (no se encontraron en el documento).
    """
    extractor = EXTRACTORES.get(tipo_documento)
    if extractor is None:
        return {}
    try:
        doc = Document(archivo)
        return extractor(doc)
    except Exception as e:
        return {'_error': str(e)}