import pdfplumber
import re
from docx import Document
from rapidfuzz import fuzz

# =========================
# DICCIONARIOS PARA CONVERSIÓN DE TEXTO A NÚMERO
# =========================

UNIDADES = {
    'cero':0,'un':1,'uno':1,'una':1,'dos':2,'tres':3,'cuatro':4,'cinco':5,
    'seis':6,'siete':7,'ocho':8,'nueve':9,'diez':10,'once':11,'doce':12,
    'trece':13,'catorce':14,'quince':15,'dieciseis':16,'diecisiete':17,
    'dieciocho':18,'diecinueve':19,'veinte':20,'veintiun':21,'veintidos':22,
    'veintitres':23,'veinticuatro':24,'veinticinco':25,'veintiseis':26,
    'veintisiete':27,'veintiocho':28,'veintinueve':29,
}
DECENAS = {
    'treinta':30,'cuarenta':40,'cincuenta':50,'sesenta':60,
    'setenta':70,'ochenta':80,'noventa':90,
}
CENTENAS = {
    'cien':100,'ciento':100,'doscientos':200,'doscientas':200,
    'trescientos':300,'trescientas':300,'cuatrocientos':400,'cuatrocientas':400,
    'quinientos':500,'quinientas':500,'seiscientos':600,'seiscientas':600,
    'setecientos':700,'setecientas':700,'ochocientos':800,'ochocientas':800,
    'novecientos':900,'novecientas':900,
}

MESES_ES = {
    'enero':'01','febrero':'02','marzo':'03','abril':'04',
    'mayo':'05','junio':'06','julio':'07','agosto':'08',
    'septiembre':'09','octubre':'10','noviembre':'11','diciembre':'12',
}

def texto_a_numero(texto):
    """
    Convierte 'CUATRO MILLONES OCHOCIENTOS CATORCE MIL SETECIENTOS TRES' → 4814703
    Retorna None si no puede interpretar.
    """
    texto = re.sub(r'pesos?\s*(colombianos?)?', '', texto, flags=re.IGNORECASE)
    texto = re.sub(r'\s+', ' ', texto).strip().lower()
    texto = texto.replace('á','a').replace('é','e').replace('í','i').replace('ó','o').replace('ú','u')

    total = 0
    millones = 0
    miles = 0
    resto = 0
    palabras = texto.split()
    acum = 0

    i = 0
    while i < len(palabras):
        p = palabras[i]
        if p in ('y',):
            i += 1
            continue
        if p == 'mill' and i+1 < len(palabras) and palabras[i+1] in ('on','ones'):
            millones = acum if acum else 1
            acum = 0
            i += 2
            continue
        if p in ('millon', 'millones'):
            millones = acum if acum else 1
            acum = 0
        elif p in ('mil',):
            miles = acum if acum else 1
            acum = 0
        elif p in CENTENAS:
            acum += CENTENAS[p]
        elif p in DECENAS:
            acum += DECENAS[p]
        elif p in UNIDADES:
            acum += UNIDADES[p]
        i += 1

    resto = acum
    total = millones * 1_000_000 + miles * 1_000 + resto
    return total if total > 0 else None


# =========================
# NORMALIZAR VALOR MONETARIO
# =========================

def normalizar_monto(valor):
    """Quita $, puntos de miles y comas → entero comparable."""
    if not valor:
        return None
    limpio = re.sub(r'[\$\s.]', '', str(valor))
    limpio = limpio.replace(',', '')
    try:
        return int(limpio)
    except ValueError:
        return None


def formatear_monto(numero):
    """Formatea entero como $ 4.151.984"""
    if numero is None:
        return "N/A"
    return f"$ {numero:,.0f}".replace(",", ".")


# =========================
# BUSCAR EN FORMATO TABLA (con normalización de acentos)
# =========================

def buscar_en_tabla(etiquetas, texto):
    """
    Busca patrones tipo:
        ETIQUETA | valor
        ETIQUETA\nvalor  (celda siguiente en DOCX)
    Retorna el valor encontrado o None.
    """
    def norm(t):
        return t.replace('É','E').replace('é','e').replace('Á','A').replace('á','a')\
                .replace('Í','I').replace('í','i').replace('Ó','O').replace('ó','o')\
                .replace('Ú','U').replace('ú','u')

    texto_norm = norm(texto)
    etiquetas_norm = [norm(e) for e in etiquetas]
    lineas = texto_norm.split('\n')

    for i, linea in enumerate(lineas):
        linea_limpia = linea.strip()
        for etiqueta in etiquetas_norm:
            if '|' in linea_limpia:
                partes = [p.strip() for p in linea_limpia.split('|')]
                for j, parte in enumerate(partes):
                    if fuzz.token_sort_ratio(limpiar(etiqueta), limpiar(parte)) >= 75:
                        if j + 1 < len(partes) and partes[j + 1]:
                            return partes[j + 1]
            if fuzz.token_sort_ratio(limpiar(etiqueta), limpiar(linea_limpia)) >= 75:
                if i + 1 < len(lineas) and lineas[i + 1].strip():
                    return lineas[i + 1].strip()

    return None


# =========================
# VALIDAR NÚMERO DE CONTACTO
# =========================

def validar_numero_contacto(etiquetas, valor_usuario, texto):
    valor_limpio = re.sub(r'\D', '', str(valor_usuario))
    lineas = texto.split('\n')

    def norm(t):
        return t.replace('É','E').replace('é','e').replace('Á','A').replace('á','a')\
                .replace('Í','I').replace('í','i').replace('Ó','O').replace('ó','o')\
                .replace('Ú','U').replace('ú','u')

    # 1. Buscar "Etiqueta: número"
    for etiqueta in etiquetas:
        patron = re.compile(
            rf'{re.escape(norm(etiqueta))}\s*[:\-]?\s*(\d[\d\s\-]{{6,14}}\d)',
            re.IGNORECASE
        )
        for match in patron.finditer(norm(texto)):
            encontrado_limpio = re.sub(r'\D', '', match.group(1))
            if abs(len(encontrado_limpio) - len(valor_limpio)) <= 2:
                coincide = encontrado_limpio == valor_limpio
                return {
                    "estado": "CORRECTO" if coincide else "DIFERENTE",
                    "encontrado": encontrado_limpio,
                    "mensaje": None if coincide else
                        f'Se ingresó "{valor_usuario}" pero el documento tiene "{encontrado_limpio}"',
                }

    # 2. Buscar en tabla pipe: etiqueta en col j, valor en col j de la siguiente fila pipe
    lineas_pipe = [(i, linea) for i, linea in enumerate(lineas) if '|' in linea]
    for idx, (i, linea) in enumerate(lineas_pipe):
        encabezados = [p.strip() for p in linea.split('|')]
        for j, enc in enumerate(encabezados):
            score_max = max(
                (fuzz.token_sort_ratio(limpiar(etiqueta), limpiar(enc)) for etiqueta in etiquetas),
                default=0
            )
            if score_max >= 75:
                for k, otra in enumerate(encabezados):
                    if k == j:
                        continue
                    num = re.sub(r'\D', '', otra)
                    if len(num) >= 7 and abs(len(num) - len(valor_limpio)) <= 2:
                        coincide = num == valor_limpio
                        return {
                            "estado": "CORRECTO" if coincide else "DIFERENTE",
                            "encontrado": num,
                            "mensaje": None if coincide else
                                f'Se ingresó "{valor_usuario}" pero el documento tiene "{num}"',
                        }
                if idx + 1 < len(lineas_pipe):
                    _, sig_linea = lineas_pipe[idx + 1]
                    sig_celdas = [p.strip() for p in sig_linea.split('|')]
                    if j < len(sig_celdas):
                        num = re.sub(r'\D', '', sig_celdas[j])
                        if len(num) >= 7 and abs(len(num) - len(valor_limpio)) <= 2:
                            coincide = num == valor_limpio
                            return {
                                "estado": "CORRECTO" if coincide else "DIFERENTE",
                                "encontrado": num,
                                "mensaje": None if coincide else
                                    f'Se ingresó "{valor_usuario}" pero el documento tiene "{num}"',
                            }

    return {
        "estado": "NO_ENCONTRADO",
        "encontrado": None,
        "mensaje": "No se encontró el número de contacto en el documento",
    }


# =========================
# VALIDAR CAMPO GENÉRICO (texto)
# =========================

def validar_campo_texto(etiquetas, valor_usuario, texto, umbral=80):
    valor_usuario = valor_usuario.strip()
    valor_limpio = limpiar(valor_usuario)

    # 1. Buscar formato inline: "Etiqueta: valor"
    for etiqueta in etiquetas:
        patron = re.compile(
            rf'{re.escape(etiqueta)}\s*[:\-]?\s*(.{{3,80}})',
            re.IGNORECASE
        )
        for match in patron.finditer(texto):
            encontrado = match.group(1).strip().split('\n')[0].strip()
            if len(encontrado.split()) > 10:
                continue
            score = fuzz.token_sort_ratio(valor_limpio, limpiar(encontrado))
            if score >= umbral:
                return {
                    "estado": "CORRECTO" if score == 100 else "SIMILAR",
                    "encontrado": encontrado,
                    "score": score,
                    "mensaje": None,
                }
            else:
                return {
                    "estado": "DIFERENTE",
                    "encontrado": encontrado,
                    "score": score,
                    "mensaje": f'Se ingresó "{valor_usuario}" pero el documento dice "{encontrado}"',
                }

    # 2. Buscar en tablas (pipe o línea siguiente)
    encontrado_tabla = buscar_en_tabla(etiquetas, texto)
    if encontrado_tabla:
        score = fuzz.token_sort_ratio(valor_limpio, limpiar(encontrado_tabla))
        if score >= umbral:
            return {
                "estado": "CORRECTO" if score == 100 else "SIMILAR",
                "encontrado": encontrado_tabla,
                "score": score,
                "mensaje": None,
            }
        else:
            return {
                "estado": "DIFERENTE",
                "encontrado": encontrado_tabla,
                "score": score,
                "mensaje": f'Se ingresó "{valor_usuario}" pero el documento dice "{encontrado_tabla}"',
            }

    return {
        "estado": "NO_ENCONTRADO",
        "encontrado": None,
        "score": 0,
        "mensaje": "No se encontró el campo en el documento",
    }


# =========================
# VALIDAR MONTO (separando número y letras)
# =========================

def validar_monto(etiquetas, valor_usuario, texto):
    monto_usuario = normalizar_monto(valor_usuario)
    if monto_usuario is None:
        return {"estado": "SIN_DATO", "encontrado": None, "mensaje": None}

    def _comparar(encontrado_raw):
        resultados_parciales = []

        # 1. Buscar número explícito ($4.814.703 o ($4.814.703))
        nums = re.findall(r'[\(\[]?\$?\s*[\d][,.\d]+[\)\]]?', encontrado_raw)
        numero_encontrado = None
        for n in nums:
            n_limpio = re.sub(r'[\(\)\[\]]', '', n)
            m = normalizar_monto(n_limpio)
            if m and m > 1000:
                numero_encontrado = m
                break

        if numero_encontrado is not None:
            if numero_encontrado == monto_usuario:
                resultados_parciales.append(("numero", "OK", None))
            else:
                resultados_parciales.append((
                    "numero", "ERROR",
                    f'Número: se ingresó {formatear_monto(monto_usuario)} pero el documento tiene {formatear_monto(numero_encontrado)}'
                ))

        # 2. Buscar texto en letras
        m_letras = texto_a_numero(encontrado_raw)
        if m_letras is not None:
            if m_letras == monto_usuario:
                resultados_parciales.append(("letras", "OK", None))
            else:
                resultados_parciales.append((
                    "letras", "ERROR",
                    "No coincide el valor en letras con el valor ingresado"
                ))

        if not resultados_parciales:
            return None

        errores = [msg for _, estado, msg in resultados_parciales if estado == "ERROR"]
        advertencias = []

        # Si hay número pero no hay letras detectables → advertir que no se pudo validar letras
        tiene_numero = any(t == "numero" for t, _, _ in resultados_parciales)
        tiene_letras = any(t == "letras" for t, _, _ in resultados_parciales)

        if tiene_numero and not tiene_letras:
            advertencias.append("⚠️ No se pudo validar el valor en letras")

        if errores:
            return {
                "estado": "DIFERENTE",
                "encontrado": encontrado_raw,
                "mensaje": " | ".join(errores),
            }
        elif advertencias:
            return {
                "estado": "CORRECTO",
                "encontrado": encontrado_raw,
                "mensaje": " | ".join(advertencias),
            }
        else:
            return {
                "estado": "CORRECTO",
                "encontrado": encontrado_raw,
                "mensaje": None,
            }

    for etiqueta in etiquetas:
        patron = re.compile(
            rf'{re.escape(etiqueta)}\s*[:\-]?\s*(.{{3,120}})',
            re.IGNORECASE
        )
        for match in patron.finditer(texto):
            encontrado = match.group(1).strip().split('\n')[0]
            r = _comparar(encontrado)
            if r:
                return r

    encontrado_tabla = buscar_en_tabla(etiquetas, texto)
    if encontrado_tabla:
        r = _comparar(encontrado_tabla)
        if r:
            return r

    return {
        "estado": "NO_ENCONTRADO",
        "encontrado": None,
        "mensaje": "No se encontró el monto en el documento",
    }


# =========================
# VALIDAR FECHA
# =========================

def normalizar_fecha(texto_fecha):
    texto_fecha = texto_fecha.strip().lower()

    m = re.match(r'(\d{1,2})\s+de\s+(\w+)\s+de\s+(\d{4})', texto_fecha)
    if m:
        dia, mes_str, anio = m.groups()
        mes = MESES_ES.get(mes_str)
        if mes:
            return f"{anio}-{mes}-{dia.zfill(2)}"

    m = re.match(r'(\d{1,2})[\/\-](\d{1,2})[\/\-](\d{4})', texto_fecha)
    if m:
        dia, mes, anio = m.groups()
        return f"{anio}-{mes.zfill(2)}-{dia.zfill(2)}"

    m = re.match(r'(\d{4})[\/\-](\d{2})[\/\-](\d{2})', texto_fecha)
    if m:
        return texto_fecha[:10]

    return texto_fecha


def validar_fecha(etiquetas, valor_usuario, texto):
    if not valor_usuario:
        return {"estado": "SIN_DATO", "encontrado": None, "mensaje": None}

    fecha_usuario_norm = normalizar_fecha(str(valor_usuario))

    patron_fecha = re.compile(
        r'(\d{1,2}\s+de\s+\w+\s+de\s+\d{4}|\d{1,2}[\/\-]\d{1,2}[\/\-]\d{4}|\d{4}[\/\-]\d{2}[\/\-]\d{2})',
        re.IGNORECASE
    )

    for etiqueta in etiquetas:
        bloque = re.compile(
            rf'{re.escape(etiqueta)}\s*[:\-]?\s*(.{{0,80}})',
            re.IGNORECASE
        )
        for match in bloque.finditer(texto):
            fragmento = match.group(1)
            fm = patron_fecha.search(fragmento)
            if fm:
                encontrado = fm.group(0).strip()
                encontrado_norm = normalizar_fecha(encontrado)
                if fecha_usuario_norm == encontrado_norm:
                    return {"estado": "CORRECTO", "encontrado": encontrado, "mensaje": None}
                else:
                    return {
                        "estado": "DIFERENTE",
                        "encontrado": encontrado,
                        "mensaje": f'Se ingresó "{valor_usuario}" pero el documento tiene "{encontrado}"',
                    }

    encontrado_tabla = buscar_en_tabla(etiquetas, texto)
    if encontrado_tabla:
        fm = patron_fecha.search(encontrado_tabla)
        if fm:
            encontrado = fm.group(0).strip()
            encontrado_norm = normalizar_fecha(encontrado)
            if fecha_usuario_norm == encontrado_norm:
                return {"estado": "CORRECTO", "encontrado": encontrado, "mensaje": None}
            else:
                return {
                    "estado": "DIFERENTE",
                    "encontrado": encontrado,
                    "mensaje": f'Se ingresó "{valor_usuario}" pero el documento tiene "{encontrado}"',
                }

    return {"estado": "NO_ENCONTRADO", "encontrado": None, "mensaje": "No se encontró la fecha en el documento"}


# =========================
# VALIDAR TODOS LOS CAMPOS
# =========================

def validar_campos_adicionales(form_data, texto):
    campos = []

    definiciones = [
        {
            "clave": "numero_contacto",
            "label": "Número de contacto",
            "etiquetas": ["TELÉFONO", "Teléfono", "Telefono", "Numero de contacto", "Celular"],
            "tipo": "contacto",
        },
        {
            "clave": "cargo",
            "label": "Cargo",
            "etiquetas": ["OFICIO QUE DESEMPEÑARA EL TRABAJADOR", "Oficio que desempeñara", "Cargo de la persona", "Cargo"],
            "tipo": "texto",
        },
        {
            "clave": "categoria_salarial",
            "label": "Categoría salarial",
            "etiquetas": ["Categoria Salarial", "Categoría Salarial"],
            "tipo": "texto",
        },
        {
            "clave": "salario_basico",
            "label": "Salario básico mensual",
            "etiquetas": ["SALARIO MENSUAL", "Salario basico mensual", "Salario básico mensual", "Salario básico"],
            "tipo": "monto",
        },
        {
            "clave": "beneficio_alimentacion",
            "label": "Beneficio alimentación",
            "etiquetas": ["Beneficios de alimentación", "Beneficio alimentacion", "Alimentación"],
            "tipo": "monto",
        },
        {
            "clave": "beneficio_pension",
            "label": "Beneficio pensión voluntaria",
            "etiquetas": ["Beneficio pension voluntaria", "Pensión voluntaria"],
            "tipo": "monto",
        },
        {
            "clave": "auxilio_localizacion",
            "label": "Auxilio de localización",
            "etiquetas": ["Auxilio de localización", "Auxilio de localizacion"],
            "tipo": "monto",
        },
        {
            "clave": "auxilio_vivienda",
            "label": "Auxilio de vivienda",
            "etiquetas": ["Auxilio de vivienda"],
            "tipo": "texto",
        },
        {
            "clave": "condicion_descanso",
            "label": "Condición de descanso",
            "etiquetas": ["Condición de descanso", "Condicion de descanso"],
            "tipo": "texto",
        },
        {
            "clave": "tipo_contrato",
            "label": "Tipo de contrato",
            "etiquetas": ["Tipo de contrato"],
            "tipo": "texto",
        },
        {
            "clave": "fecha_inicio",
            "label": "Fecha de inicio",
            "etiquetas": ["FECHA DE INICIO DE LABORES", "Fecha de inicio de labores", "Fecha de inicio"],
            "tipo": "fecha",
        },
        {
            "clave": "proyecto",
            "label": "Proyecto / Centro de costos",
            "etiquetas": ["Proyecto / centro de costos", "Proyecto", "Centro de costos"],
            "tipo": "texto",
        },
    ]

    for defn in definiciones:
        valor = form_data.get(defn["clave"])
        if not valor:
            continue

        if defn["tipo"] == "monto":
            resultado = validar_monto(defn["etiquetas"], valor, texto)
        elif defn["tipo"] == "fecha":
            resultado = validar_fecha(defn["etiquetas"], valor, texto)
        elif defn["tipo"] == "contacto":
            resultado = validar_numero_contacto(defn["etiquetas"], valor, texto)
        else:
            resultado = validar_campo_texto(defn["etiquetas"], valor, texto)

        campos.append({
            "label": defn["label"],
            "valor_usuario": str(valor),
            "tipo": defn["tipo"],
            **resultado,
        })

    return campos


# =========================
# EXTRAER TEXTO (PDF mejorado)
# =========================

def extraer_texto_pdf(archivo):
    texto = ""
    try:
        with pdfplumber.open(archivo) as pdf:
            for pagina in pdf.pages:
                palabras = pagina.extract_words(
                    keep_blank_chars=True,
                    use_text_flow=True,
                    extra_attrs=["fontname", "size"]
                )
                if palabras:
                    palabras_ordenadas = sorted(palabras, key=lambda w: (w["top"], w["x0"]))
                    linea_actual = ""
                    ultimo_top = None
                    ultimo_x1 = 0
                    for w in palabras_ordenadas:
                        if ultimo_top is None or abs(w["top"] - ultimo_top) > 5:
                            if linea_actual:
                                texto += linea_actual.strip() + "\n"
                            linea_actual = w["text"]
                            ultimo_top = w["top"]
                            ultimo_x1 = w["x1"]
                        else:
                            espacio = w["x0"] - ultimo_x1
                            if espacio > 3:
                                linea_actual += " " + w["text"]
                            else:
                                linea_actual += w["text"]
                            ultimo_x1 = w["x1"]
                    if linea_actual:
                        texto += linea_actual.strip() + "\n"
                else:
                    contenido = pagina.extract_text()
                    if contenido:
                        texto += contenido + "\n"

                if pagina.annots:
                    for ann in pagina.annots:
                        if ann.get("contents"):
                            texto += ann["contents"] + "\n"
                        if ann.get("title"):
                            texto += ann["title"] + "\n"

    except Exception as e:
        print(f"Error extrayendo PDF: {e}")
    return texto


def extraer_texto_docx(archivo):
    texto = ""
    try:
        doc = Document(archivo)
        for p in doc.paragraphs:
            if p.text.strip():
                texto += p.text + "\n"
        for tabla in doc.tables:
            filas = list(tabla.rows)
            n_filas = len(filas)
            for i, fila in enumerate(filas):
                celdas = [c.text.strip() for c in fila.cells if c.text.strip()]
                if not celdas:
                    continue

                texto += " | ".join(celdas) + "\n"

                if i + 1 < len(filas):
                    sig_celdas = [c.text.strip() for c in filas[i + 1].cells if c.text.strip()]
                    encabezados_mayus = all(c == c.upper() for c in celdas)
                    valores_no_mayus = any(c != c.upper() or any(ch.isdigit() for ch in c) for c in sig_celdas)

                    if encabezados_mayus and sig_celdas and len(celdas) == len(sig_celdas):
                        for enc, val in zip(celdas, sig_celdas):
                            texto += f"{enc}: {val}\n"

    except Exception as e:
        print(f"Error extrayendo DOCX: {e}")
    return texto


def extraer_texto(archivo):
    nombre = archivo.name.lower()
    if nombre.endswith(".pdf"):
        return extraer_texto_pdf(archivo)
    elif nombre.endswith(".docx"):
        return extraer_texto_docx(archivo)
    return ""


# =========================
# LIMPIEZA
# =========================

def limpiar(texto):
    texto = texto.upper()
    texto = re.sub(r'[^A-Z0-9\s]', ' ', texto)
    texto = re.sub(r'\s+', ' ', texto)
    return texto.strip()


# =========================
# BUSCAR TODAS LAS COINCIDENCIAS DEL NOMBRE
# =========================

def encontrar_ocurrencias_nombre(nombre_usuario, texto, umbral=80):
    texto_limpio = limpiar(texto)
    palabras = texto_limpio.split()
    cantidad_palabras = len(nombre_usuario.split())
    ocurrencias = []

    for i in range(len(palabras)):
        for extra in range(3):
            fragmento = " ".join(palabras[i:i + cantidad_palabras + extra])
            if not fragmento:
                continue
            score = fuzz.token_sort_ratio(nombre_usuario, fragmento)
            if score >= umbral:
                ocurrencias.append({
                    "texto": fragmento,
                    "score": score,
                    "inicio": i
                })
                break

    ocurrencias_filtradas = []
    for occ in sorted(ocurrencias, key=lambda x: x["inicio"]):
        if not ocurrencias_filtradas:
            ocurrencias_filtradas.append(occ)
        else:
            anterior = ocurrencias_filtradas[-1]
            if occ["inicio"] < anterior["inicio"] + cantidad_palabras + 2:
                if occ["score"] > anterior["score"]:
                    ocurrencias_filtradas[-1] = occ
            else:
                ocurrencias_filtradas.append(occ)

    return ocurrencias_filtradas


# =========================
# VALIDAR NOMBRE
# =========================

def validar_nombre(nombre_usuario, texto, umbral=80):
    nombre_usuario = limpiar(nombre_usuario)
    ocurrencias = encontrar_ocurrencias_nombre(nombre_usuario, texto, umbral)

    total = len(ocurrencias)
    exactas = 0
    con_error = 0
    detalles_ocurrencias = []

    for occ in ocurrencias:
        palabras_usuario = nombre_usuario.split()
        palabras_occ = occ["texto"].split()
        detalles_palabras = []
        for palabra in palabras_usuario:
            mejor = None
            mejor_score = 0
            for p in palabras_occ:
                score = fuzz.ratio(palabra, p)
                if score > mejor_score:
                    mejor_score = score
                    mejor = p
            if mejor_score == 100:
                detalles_palabras.append({"estado": "OK", "palabra_usuario": palabra, "encontrado": mejor})
            elif mejor_score >= 80:
                detalles_palabras.append({"estado": "ERROR", "palabra_usuario": palabra, "encontrado": mejor})
            else:
                detalles_palabras.append({"estado": "FALTA", "palabra_usuario": palabra, "encontrado": None})

        adicionales = [p for p in palabras_occ if p not in palabras_usuario]
        es_exacta = all(d["estado"] == "OK" for d in detalles_palabras) and not adicionales

        if es_exacta:
            exactas += 1
        else:
            con_error += 1

        detalles_ocurrencias.append({
            "texto": occ["texto"],
            "score": occ["score"],
            "es_exacta": es_exacta,
            "detalles": detalles_palabras,
            "adicionales": adicionales,
        })

    promedio = sum(o["score"] for o in ocurrencias) / total if total > 0 else 0

    return {
        "nombre_usuario": nombre_usuario,
        "total_ocurrencias": total,
        "exactas": exactas,
        "con_errores": con_error,
        "porcentaje_promedio": round(promedio, 2),
        "ocurrencias": detalles_ocurrencias,
    }


# =========================
# DETECTAR TIPO DE DOCUMENTO
# =========================

def detectar_tipo_documento(documento):
    documento_limpio = re.sub(r'\D', '', str(documento))
    longitud = len(documento_limpio)

    if longitud == 9 or longitud == 10:
        if longitud == 9:
            return {
                "tipo": "NIT",
                "descripcion": "Número de Identificación Tributaria",
                "longitud_esperada": 9,
                "patron": "\\d{9}"
            }
        else:
            return {
                "tipo": "NIT_O_CC",
                "descripcion": "NIT (10 dígitos) o Cédula (10 dígitos)",
                "longitud_esperada": 10,
                "patron": "\\d{10}"
            }
    elif 6 <= longitud <= 8:
        return {
            "tipo": "CC",
            "descripcion": "Cédula de Ciudadanía",
            "longitud_esperada": longitud,
            "patron": f"\\d{{{longitud}}}"
        }
    elif 11 <= longitud <= 12:
        return {
            "tipo": "CE",
            "descripcion": "Cédula de Extranjería",
            "longitud_esperada": longitud,
            "patron": f"\\d{{{longitud}}}"
        }
    else:
        return {
            "tipo": "DESCONOCIDO",
            "descripcion": f"Documento de {longitud} dígitos",
            "longitud_esperada": longitud,
            "patron": f"\\d{{{longitud}}}"
        }


def validar_nit_colombia(nit):
    nit = re.sub(r'\D', '', str(nit))
    if len(nit) not in [9, 10]:
        return {"valido": False, "razon": f"NIT debe tener 9 o 10 dígitos. Tiene {len(nit)}"}

    if len(nit) == 10:
        pesos = [3, 7, 13, 17, 19, 23, 29, 31, 37]
        suma = sum(int(nit[i]) * pesos[i] for i in range(9))
        dv_calc = suma % 11
        if dv_calc == 0:
            dv_calc = 0
        elif dv_calc == 1:
            dv_calc = 9
        else:
            dv_calc = 11 - dv_calc
        dv_real = int(nit[9])
        if dv_real == dv_calc:
            return {"valido": True, "razon": f"NIT válido. Dígito verificador correcto: {dv_real}"}
        else:
            return {"valido": False, "razon": f"NIT inválido. Dígito verificador debería ser {dv_calc}, pero es {dv_real}"}

    return {"valido": True, "razon": "NIT de 9 dígitos (sin verificador). No se puede validar completamente."}


# =========================
# Describir diferencia entre dos números (versión simple)
# =========================

def describir_diferencia_numeros(usuario, encontrado):
    u = re.sub(r'\D', '', str(usuario))
    e = re.sub(r'\D', '', str(encontrado))

    if len(u) != len(e):
        diff = len(e) - len(u)
        if diff > 0:
            return f"Tiene {diff} dígito(s) de más"
        else:
            return f"Tiene {abs(diff)} dígito(s) de menos"

    posiciones = [i + 1 for i in range(len(u)) if u[i] != e[i]]
    if posiciones:
        return f"No coincide con el número ingresado"

    return "Coincide exactamente"


# =========================
# Patrón CC mejorado — captura todas las variantes
# =========================

PATRON_CC = re.compile(
    r'(C\.?\s*C\.?|C[eéEÉ]dula(?:\s+[Dd]e\s+[Cc]iudadan[ií]a)?)\s*[:\-]?\s*(\d[\d.\s]{4,}\d|\d{5,})',
    re.IGNORECASE
)


def detectar_contextos_cc(texto, documento_usuario):
    contextos = []
    for match in PATRON_CC.finditer(texto):
        etiqueta_raw = match.group(1).strip()
        numero_raw = match.group(2)
        numero_limpio = re.sub(r'\D', '', numero_raw)

        if not numero_limpio:
            continue

        etiqueta_upper = etiqueta_raw.upper().replace(' ', '')
        if re.match(r'C\.?C\.?', etiqueta_upper):
            etiqueta_display = "C.C."
        elif 'CEDULA' in etiqueta_upper or 'CÉDULA' in etiqueta_upper:
            if 'CIUDADAN' in etiqueta_upper:
                etiqueta_display = "Cédula de Ciudadanía"
            else:
                etiqueta_display = "Cédula"
        else:
            etiqueta_display = etiqueta_raw

        start = match.start()
        end = match.end()
        antes = texto[max(0, start - 40):start].strip()
        despues = texto[end:end + 40].strip()

        coincide = (numero_limpio == re.sub(r'\D', '', str(documento_usuario)))
        diferencia = None if coincide else describir_diferencia_numeros(documento_usuario, numero_limpio)

        contextos.append({
            "etiqueta": etiqueta_display,
            "numero_encontrado": numero_limpio,
            "antes": antes,
            "despues": despues,
            "coincide": coincide,
            "diferencia": diferencia,
        })

    return contextos


# =========================
# VALIDAR DOCUMENTO
# =========================

def validar_documento(documento_usuario, texto, umbral=80, longitud_minima=5,
                      numeros_ignorados=None):
    try:
        documento_usuario = str(documento_usuario).strip()
        doc_solo_digitos = re.sub(r'\D', '', documento_usuario)
        longitud_usuario = len(doc_solo_digitos)

        numeros_ignorados = set(numeros_ignorados or [])

        tipo_info = detectar_tipo_documento(documento_usuario)
        tipo_documento = tipo_info["tipo"]
        descripcion_tipo = tipo_info["descripcion"]

        todos_numeros = re.findall(r'\d+', texto)

        TOLERANCIA = 1
        numeros_relevantes = [
            n for n in todos_numeros
            if abs(len(n) - longitud_usuario) <= TOLERANCIA
            and len(n) >= longitud_minima
            and n not in numeros_ignorados
        ]
        numeros_exactos = [
            n for n in todos_numeros
            if n == doc_solo_digitos and n not in numeros_ignorados
        ]
        numeros_relevantes = list(dict.fromkeys(numeros_relevantes + numeros_exactos))

        ocurrencias = []
        coincidencias_exactas = []

        for num in numeros_relevantes:
            sim = fuzz.ratio(doc_solo_digitos, num)
            es_exacto = (num == doc_solo_digitos)

            diferencia_detalle = None if es_exacto else describir_diferencia_numeros(doc_solo_digitos, num)

            ocurrencias.append({
                "numero": num,
                "similitud": sim,
                "es_exacto": es_exacto,
                "largo": len(num),
                "diferencia_detalle": diferencia_detalle,
            })

            if es_exacto:
                coincidencias_exactas.append(num)

        total = len(ocurrencias)
        exactas = len(coincidencias_exactas)
        con_error = total - exactas

        numeros_unicos = set(o["numero"] for o in ocurrencias) if ocurrencias else set()
        inconsistencia = len(numeros_unicos) > 1

        validacion_nit = None
        if tipo_documento == "NIT" and exactas > 0:
            nit_encontrado = list(coincidencias_exactas)[0]
            validacion_nit = validar_nit_colombia(nit_encontrado)

        contextos_cc = detectar_contextos_cc(texto, documento_usuario)

        cc_exacto = any(c["coincide"] for c in contextos_cc)

        if len(contextos_cc) > 0:
            if cc_exacto:
                estado = "CORRECTO"
                mensaje = f"✅ CÉDULA CORRECTA: '{documento_usuario}' ({descripcion_tipo}) encontrado junto a etiqueta de cédula."
            else:
                nums_cc = ", ".join(
                    f"{c['numero_encontrado']} ({c['diferencia']})"
                    for c in contextos_cc if c["numero_encontrado"]
                )
                estado = "DOCUMENTO_INCORRECTO"
                mensaje = (
                    f"❌ CÉDULA INCORRECTA: Se ingresó '{documento_usuario}' pero en el documento "
                    f"la cédula registrada es diferente. "
                    f"Encontrado: {nums_cc}"
                )
        elif total == 0:
            estado = "NO_ENCONTRADO"
            mensaje = (
                f"⚠️ CÉDULA NO ENCONTRADA: Se ingresó '{documento_usuario}' ({descripcion_tipo}) "
                f"pero NO se encontró en el documento."
            )
        elif exactas > 0 and con_error == 0:
            estado = "CORRECTO"
            mensaje = f"✅ CÉDULA CORRECTA: '{documento_usuario}' ({descripcion_tipo}) en {exactas} ocurrencia(s)."
            if validacion_nit and not validacion_nit["valido"]:
                mensaje += f"\n⚠️ ADVERTENCIA NIT: {validacion_nit['razon']}"
                estado = "CORRECTO_PERO_NIT_INVALIDO"
        elif inconsistencia and exactas == 0:
            numeros_encontrados = ", ".join(sorted(numeros_unicos))
            estado = "DOCUMENTO_INCORRECTO"
            mensaje = (
                f"❌ ERROR CRÍTICO: Se ingresó '{documento_usuario}' ({descripcion_tipo}) "
                f"pero se encontraron números DIFERENTES: {numeros_encontrados}. VERIFICAR DOCUMENTO."
            )
        elif inconsistencia and exactas > 0:
            numeros_str = ", ".join(sorted(numeros_unicos))
            estado = "INCONSISTENCIA"
            mensaje = (
                f"⚠️ INCONSISTENCIA: Se encontraron {total} versiones del documento. "
                f"{exactas} coinciden exactamente, {con_error} con diferencias. "
                f"Versiones: {numeros_str}"
            )
        else:
            similitudes_str = ", ".join([f"{o['numero']} ({o['similitud']}%)" for o in ocurrencias])
            estado = "PARCIAL"
            mensaje = (
                f"⚠️ CÉDULA PARCIAL: Se encontraron {total} número(s) similares "
                f"pero NINGUNO es exacto. Similitudes: {similitudes_str}"
            )

        return {
            "documento_usuario": documento_usuario,
            "tipo_documento": tipo_documento,
            "descripcion_tipo": descripcion_tipo,
            "total_ocurrencias": total,
            "exactas": exactas,
            "con_errores": con_error,
            "inconsistencia_entre_ocurrencias": inconsistencia,
            "estado": estado,
            "ocurrencias": ocurrencias,
            "mensaje": mensaje,
            "todos_numeros_documento": todos_numeros,
            "validacion_nit": validacion_nit,
            "contextos_cc": contextos_cc,
        }

    except Exception as e:
        return {
            "documento_usuario": str(documento_usuario),
            "tipo_documento": "ERROR",
            "descripcion_tipo": "No se pudo determinar",
            "total_ocurrencias": 0,
            "exactas": 0,
            "con_errores": 0,
            "inconsistencia_entre_ocurrencias": False,
            "estado": "ERROR",
            "ocurrencias": [],
            "mensaje": f"❌ ERROR EN VALIDACIÓN: {str(e)}",
            "todos_numeros_documento": [],
            "validacion_nit": None,
            "contextos_cc": [],
            "error_detalle": str(e),
        }


# =========================
# DEFINICIÓN DE CAMPOS POR TIPO DE DOCUMENTO
# =========================

CAMPOS_POR_TIPO = {
    "AUXILIO": [
        "nombre",
        "documento",
        "cargo",
        "proyecto",
        "fecha_inicio",
        "auxilio_localizacion",
        "ciudad",
        "aux_almuerzo",
        "aux_transporte",
        "aux_vivienda",
        "beneficio_alimentacion",
        "beneficio_pension",
        "aux_desplazamiento",
    ],
    "CONDICIONES": [
        "nombre",
        "cargo",
        "proyecto",
        "fecha_inicio",
        "salario_basico",
        "beneficio_alimentacion",
        "auxilio_localizacion",
        "beneficio_pension",
        "condicion_descanso",
    ],
    "CONTRATO": [
        "nombre",
        "documento",
        "cargo",
        "fecha_inicio",
        "salario_basico",
    ],
    "MAN_CONF": [
        "nombre",
        "documento",
        "fecha_inicio",
    ],
    "MEDICAS": [
        "cargo",
        "nombre",
        "fecha_inicio",
    ],
    "RIESGO": [
        "nombre",
        "documento",
        "cargo",
        "fecha_inicio",
        "proyecto",
        "tipo_riesgo",
    ],
}


# =========================
# DEFINICIÓN DETALLADA DE CAMPOS
# =========================

DEFINICIONES_CAMPOS = {
    "nombre": {
        "label": "Nombre",
        "tipo": "nombre",
    },
    "documento": {
        "label": "Documento",
        "tipo": "documento",
    },
    "numero_contacto": {
        "label": "Número de contacto",
        "etiquetas": ["TELÉFONO", "Teléfono", "Telefono", "Numero de contacto", "Celular"],
        "tipo": "contacto",
    },
    "cargo": {
        "label": "Cargo",
        "etiquetas": ["OFICIO QUE DESEMPEÑARA EL TRABAJADOR", "Oficio que desempeñara", "Cargo de la persona", "Cargo"],
        "tipo": "texto",
    },
    "categoria_salarial": {
        "label": "Categoría salarial",
        "etiquetas": ["Categoria Salarial", "Categoría Salarial"],
        "tipo": "texto",
    },
    "salario_basico": {
        "label": "Salario básico mensual",
        "etiquetas": ["SALARIO MENSUAL", "Salario basico mensual", "Salario básico mensual", "Salario básico"],
        "tipo": "monto",
    },
    "beneficio_alimentacion": {
        "label": "Beneficio alimentación",
        "etiquetas": ["Beneficios de alimentación", "Beneficio alimentacion", "Alimentación"],
        "tipo": "monto",
    },
    "beneficio_pension": {
        "label": "Beneficio pensión voluntaria",
        "etiquetas": ["Beneficio pension voluntaria", "Pensión voluntaria"],
        "tipo": "monto",
    },
    "auxilio_localizacion": {
        "label": "Auxilio de localización",
        "etiquetas": ["Auxilio de localización", "Auxilio de localizacion"],
        "tipo": "monto",
    },
    "auxilio_vivienda": {
        "label": "Auxilio de vivienda",
        "etiquetas": ["Auxilio de vivienda"],
        "tipo": "texto",
    },
    "condicion_descanso": {
        "label": "Condición de descanso",
        "etiquetas": ["Condición de descanso", "Condicion de descanso", "Descanso"],
        "tipo": "texto",
    },
    "tipo_contrato": {
        "label": "Tipo de contrato",
        "etiquetas": ["Tipo de contrato"],
        "tipo": "texto",
    },
    "fecha_inicio": {
        "label": "Fecha de inicio",
        "etiquetas": ["FECHA DE INICIO DE LABORES", "Fecha de inicio de labores", "Fecha de inicio"],
        "tipo": "fecha",
    },
    "proyecto": {
        "label": "Proyecto / Centro de costos",
        "etiquetas": ["Proyecto / centro de costos", "Proyecto", "Centro de costos"],
        "tipo": "texto",
    },
    "ciudad": {
        "label": "Ciudad de trabajo",
        "etiquetas": ["ciudad", "será en"],
        "tipo": "texto",
    },
    "aux_almuerzo": {
        "label": "Auxilio alimentación en obra (diario)",
        "etiquetas": ["AUXILIO HABITUAL Y EXTRALEGAL DE ALIMENTACIÓN", "valor diario", "AUX_ALMUERZO"],
        "tipo": "monto",
    },
    "aux_transporte": {
        "label": "Auxilio transporte en obra (diario)",
        "etiquetas": ["AUXILIO HABITUAL Y EXTRALEGAL DE TRANSPORTE", "AUX_TRANSPROTE"],
        "tipo": "monto",
    },
    "aux_vivienda": {
        "label": "Auxilio vivienda (mensual)",
        "etiquetas": ["AUXILIO EXTRALEGAL DE VIVIENDA", "AUX_VIVIENDA"],
        "tipo": "monto",
    },
    "aux_desplazamiento": {"label": "Auxilio de Desplazamiento", "tipo": "monto"},

    "tipo_riesgo": {
        "label": "Tipo de riesgo",
        "tipo": "informativo",
    },
}


def obtener_campos_por_tipo(tipo_documento):
    """
    Retorna la lista de claves de campos que deben validarse para un tipo.

    Args:
        tipo_documento: String con el tipo (AUXILIO, CONDICIONES, CONTRATO, etc.)

    Returns:
        Lista de nombres de campos a validar para ese tipo
    """
    return CAMPOS_POR_TIPO.get(tipo_documento, [])


def validar_campos_por_tipo(tipo_documento, form_data, texto):
    """
    Valida SOLO los campos específicos del tipo de documento.

    Este es el reemplazo mejorado de validar_campos_adicionales().
    En lugar de validar TODOS los campos, valida solo los aplicables
    al tipo de documento especificado.

    Args:
        tipo_documento: Tipo del documento (AUXILIO, CONDICIONES, CONTRATO, MAN_CONF, MEDICAS, RIESGO)
        form_data: cleaned_data del formulario con todos los campos
        texto: Texto extraído del archivo PDF/DOCX

    Returns:
        Lista de diccionarios con resultados de validación:
        [
            {
                "label": "Nombre del campo",
                "valor_usuario": "valor ingresado",
                "tipo": "monto|fecha|texto|contacto",
                "estado": "CORRECTO|DIFERENTE|NO_ENCONTRADO|SIN_DATO",
                "encontrado": "valor encontrado en documento",
                "mensaje": "mensaje de advertencia o error",
                ...
            },
            ...
        ]
    """
    campos_aplicables = obtener_campos_por_tipo(tipo_documento)
    campos = []

    for clave_campo in campos_aplicables:
        defn = DEFINICIONES_CAMPOS.get(clave_campo, {})
        valor = form_data.get(clave_campo)

        # Saltar campos vacíos (no se validan si no están completos)
        if not valor:
            continue

        label = defn.get("label", clave_campo)
        tipo_campo = defn.get("tipo", "texto")
        etiquetas = defn.get("etiquetas", [])

        # Los campos de nombre y documento se validan globalmente, no aquí
        if tipo_campo == "nombre":
            continue
        elif tipo_campo == "documento":
            continue

        # Validar según el tipo de campo
        elif tipo_campo == "monto":
            resultado = validar_monto(etiquetas, valor, texto)
        elif tipo_campo == "fecha":
            resultado = validar_fecha(etiquetas, valor, texto)
        elif tipo_campo == "contacto":
            resultado = validar_numero_contacto(etiquetas, valor, texto)
        else:
            # tipo_campo == "texto"
            resultado = validar_campo_texto(etiquetas, valor, texto)

        # Agregar el resultado
        campos.append({
            "label": label,
            "valor_usuario": str(valor),
            "tipo": tipo_campo,
            **resultado,
        })

    return campos


# =========================
# MEJORAS EN LA BÚSQUEDA DE CAMPOS (VERSIÓN CORREGIDA)
# =========================

def buscar_en_parrafos_y_tablas(etiquetas, texto):
    """
    Busca un valor en PÁRRAFOS Y TABLAS de manera más flexible.

    Busca patrones como:
    - "Etiqueta: valor"
    - "Etiqueta : valor"
    - "Etiqueta- valor"
    - "Etiqueta valor" (en la siguiente línea)

    Returns:
        valor encontrado o None
    """
    def norm(t):
        return t.replace('É','E').replace('é','e').replace('Á','A').replace('á','a')\
                .replace('Í','I').replace('í','i').replace('Ó','O').replace('ó','o')\
                .replace('Ú','U').replace('ú','u').strip()

    texto_norm = norm(texto)
    lineas = texto.split('\n')

    for i, linea in enumerate(lineas):
        linea_limpia = linea.strip()

        # 1. Buscar patrón "Etiqueta: valor" en esta línea
        for etiqueta in etiquetas:
            etiqueta_norm = norm(etiqueta)

            # Patrones posibles: "Etiqueta: valor" o "Etiqueta- valor" o "Etiqueta    valor"
            patrones = [
                rf'{re.escape(etiqueta_norm)}\s*:\s*(.{{1,100}})',  # Etiqueta: valor
                rf'{re.escape(etiqueta_norm)}\s*-\s*(.{{1,100}})',  # Etiqueta- valor
                rf'{re.escape(etiqueta_norm)}\s{{2,}}(.{{1,100}})',  # Etiqueta    valor (múltiples espacios) - CORREGIDO
            ]

            for patron_str in patrones:
                try:
                    patron = re.compile(patron_str, re.IGNORECASE)
                    match = patron.search(norm(linea_limpia))
                    if match:
                        valor = match.group(1).strip()
                        if valor and len(valor) < 100:  # Evitar líneas muy largas
                            return valor
                except re.error:
                    # Si hay error en regex, continuar al siguiente patrón
                    continue

        # 2. Búsqueda en tabla: "etiqueta" en col j, valor en col j siguiente
        if '|' in linea_limpia:
            celdas = [c.strip() for c in linea_limpia.split('|')]
            for j, celda in enumerate(celdas):
                celda_norm = norm(celda)
                for etiqueta in etiquetas:
                    etiqueta_norm = norm(etiqueta)
                    if fuzz.token_sort_ratio(celda_norm, etiqueta_norm) >= 75:
                        # Valor está en la siguiente celda de esta fila
                        if j + 1 < len(celdas) and celdas[j + 1]:
                            return celdas[j + 1]
                        # O está en la misma columna de la siguiente fila
                        if i + 1 < len(lineas):
                            siguiente = lineas[i + 1].split('|')
                            if j < len(siguiente) and siguiente[j].strip():
                                return siguiente[j].strip()

    return None


def validar_campo_texto_mejorado(etiquetas, valor_usuario, texto, umbral=80):
    """
    Versión mejorada de validar_campo_texto que busca en párrafos Y tablas.

    Busca el valor en:
    1. Párrafos con formato "Etiqueta: valor"
    2. Tablas con pipes
    3. Linea siguiente a la etiqueta
    """
    valor_usuario = valor_usuario.strip()
    valor_limpio = limpiar(valor_usuario)

    # 1. Buscar en párrafos y tablas con la nueva función
    encontrado_texto = buscar_en_parrafos_y_tablas(etiquetas, texto)

    if encontrado_texto:
        score = fuzz.token_sort_ratio(valor_limpio, limpiar(encontrado_texto))
        if score >= umbral:
            return {
                "estado": "CORRECTO" if score == 100 else "SIMILAR",
                "encontrado": encontrado_texto,
                "score": score,
                "mensaje": None,
            }
        else:
            return {
                "estado": "DIFERENTE",
                "encontrado": encontrado_texto,
                "score": score,
                "mensaje": f'Se ingresó "{valor_usuario}" pero el documento dice "{encontrado_texto}"',
            }

    return {
        "estado": "NO_ENCONTRADO",
        "encontrado": None,
        "score": 0,
        "mensaje": "No se encontró el campo en el documento",
    }


def validar_fecha_mejorado(etiquetas, valor_usuario, texto):
    """
    Versión mejorada de validar_fecha que busca en toda la línea, no solo después de etiqueta.

    Busca patrones de fecha en:
    1. Líneas que contienen etiqueta + fecha
    2. Tablas con fechas
    3. Línea siguiente a la etiqueta
    """
    if not valor_usuario:
        return {"estado": "SIN_DATO", "encontrado": None, "mensaje": None}

    fecha_usuario_norm = normalizar_fecha(str(valor_usuario))

    # Patrón más flexible para fechas
    patron_fecha = re.compile(
        r'(\d{1,2}\s+de\s+\w+\s+de\s+\d{4}|\d{1,2}[\/\-]\d{1,2}[\/\-]\d{4}|\d{4}[\/\-]\d{2}[\/\-]\d{2})',
        re.IGNORECASE
    )

    lineas = texto.split('\n')

    # 1. Buscar en líneas que contengan etiqueta
    for i, linea in enumerate(lineas):
        # ¿Contiene alguna etiqueta?
        for etiqueta in etiquetas:
            if fuzz.token_sort_ratio(limpiar(etiqueta), limpiar(linea[:50])) >= 60:
                # Buscar fecha en esta línea y las siguientes 2
                for offset in range(3):
                    if i + offset < len(lineas):
                        fragmento = lineas[i + offset]
                        fm = patron_fecha.search(fragmento)
                        if fm:
                            encontrado = fm.group(0).strip()
                            encontrado_norm = normalizar_fecha(encontrado)
                            if fecha_usuario_norm == encontrado_norm:
                                return {"estado": "CORRECTO", "encontrado": encontrado, "mensaje": None}
                            else:
                                return {
                                    "estado": "DIFERENTE",
                                    "encontrado": encontrado,
                                    "mensaje": f'Se ingresó "{valor_usuario}" pero el documento tiene "{encontrado}"',
                                }

    # 2. Buscar cualquier fecha en el documento (para RIESGO que busca en tabla)
    todas_fechas = patron_fecha.findall(texto)
    if todas_fechas:
        for fecha_encontrada in todas_fechas:
            encontrado_norm = normalizar_fecha(fecha_encontrada)
            if fecha_usuario_norm == encontrado_norm:
                return {"estado": "CORRECTO", "encontrado": fecha_encontrada, "mensaje": None}

    return {"estado": "NO_ENCONTRADO", "encontrado": None, "mensaje": "No se encontró la fecha en el documento"}


def validar_monto_mejorado(etiquetas, valor_usuario, texto):
    """
    Versión mejorada de validar_monto que busca en párrafos con ":" primero.
    """
    monto_usuario = normalizar_monto(valor_usuario)
    if monto_usuario is None:
        return {"estado": "SIN_DATO", "encontrado": None, "mensaje": None}

    def _comparar(encontrado_raw):
        resultados_parciales = []

        # 1. Buscar número explícito
        nums = re.findall(r'[\(\[]?\$?\s*[\d][,.\d]+[\)\]]?', encontrado_raw)
        numero_encontrado = None
        for n in nums:
            n_limpio = re.sub(r'[\(\)\[\]]', '', n)
            m = normalizar_monto(n_limpio)
            if m and m > 100:  # Monto mínimo 100
                numero_encontrado = m
                break

        if numero_encontrado is not None:
            if numero_encontrado == monto_usuario:
                resultados_parciales.append(("numero", "OK", None))
            else:
                resultados_parciales.append((
                    "numero", "ERROR",
                    f'Número: se ingresó {formatear_monto(monto_usuario)} pero el documento tiene {formatear_monto(numero_encontrado)}'
                ))

        # 2. Buscar texto en letras
        m_letras = texto_a_numero(encontrado_raw)
        if m_letras is not None:
            if m_letras == monto_usuario:
                resultados_parciales.append(("letras", "OK", None))
            else:
                resultados_parciales.append((
                    "letras", "ERROR",
                    "No coincide el valor en letras con el valor ingresado"
                ))

        if not resultados_parciales:
            return None

        errores = [msg for _, estado, msg in resultados_parciales if estado == "ERROR"]
        advertencias = []

        tiene_numero = any(t == "numero" for t, _, _ in resultados_parciales)
        tiene_letras = any(t == "letras" for t, _, _ in resultados_parciales)

        if tiene_numero and not tiene_letras:
            advertencias.append("⚠️ No se pudo validar el valor en letras")

        if errores:
            return {
                "estado": "DIFERENTE",
                "encontrado": encontrado_raw,
                "mensaje": " | ".join(errores),
            }
        elif advertencias:
            return {
                "estado": "CORRECTO",
                "encontrado": encontrado_raw,
                "mensaje": " | ".join(advertencias),
            }
        else:
            return {
                "estado": "CORRECTO",
                "encontrado": encontrado_raw,
                "mensaje": None,
            }

    # Buscar con la función mejorada
    encontrado_texto = buscar_en_parrafos_y_tablas(etiquetas, texto)
    if encontrado_texto:
        return _comparar(encontrado_texto)

    return {
        "estado": "NO_ENCONTRADO",
        "encontrado": None,
        "mensaje": "No se encontró el monto en el documento",
    }